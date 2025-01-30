import os
import csv
import re
import html
import qrcode
import json  # Import for safer JSON handling
from flask import Flask, request, jsonify, send_file, render_template
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from PIL import Image
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch, mm
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph
from reportlab.lib.utils import ImageReader
from datetime import datetime

app = Flask(__name__)

# Folder paths for file uploads, PDF outputs, and QR codes
UPLOAD_FOLDER = "uploads"
DOC_FOLDER = "label_files"
QR_FOLDER = "qr_codes"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(DOC_FOLDER, exist_ok=True)
os.makedirs(QR_FOLDER, exist_ok=True)

def load_label_dimensions(json_path):
    """
    Load and process label dimensions from a JSON file.

    Args:
        json_path (str): Path to the JSON file.

    Returns:
        dict: Dictionary with processed label dimensions, including page dimensions in inches and millimeters.

    Raises:
        Exception: If an error occurs during processing.
    """
    try:
        # Open the JSON file with UTF-8 encoding for compatibility
        with open(json_path, 'r', encoding='utf-8') as file:
            data = json.load(file)

        # Process the label dimensions
        for label_id, details in data.items():
            # Validate required keys
            required_keys = ["label_width", "label_height", "labels_per_row", "labels_per_column", "page_width_inch", "page_height_inch"]
            if not all(key in details for key in required_keys):
                raise KeyError(f"Missing required keys in label definition for {label_id}: {details}")

            # Debugging logs for dimensions
            print(f"Processing label {label_id}: {details}")

        return data

    except FileNotFoundError:
        raise Exception(f"JSON file not found: {json_path}")
    except json.JSONDecodeError as e:
        raise Exception(f"Invalid JSON format in file {json_path}: {e}")
    except KeyError as e:
        raise Exception(f"Key error while processing labels: {e}")
    except Exception as e:
        raise Exception(f"An unexpected error occurred while loading label dimensions: {e}")
    
# Base directory where the script resides
BASE_DIR = os.path.abspath(os.path.dirname(__file__))

# Construct the path to labels_types.json
LABEL_JSON_PATH = os.path.join(BASE_DIR, "templates", "labels_types.json")

# Load label dimensions
LABEL_DIMENSIONS = load_label_dimensions(LABEL_JSON_PATH)


def validate_json(json_string):
    """
    Validate the JSON structure and print its content.

    Args:
        json_string (str): The JSON string to validate.

    Returns:
        dict: Parsed JSON if valid.
    Raises:
        ValueError: If the JSON is invalid.
    """
    try:
        data = json.loads(json_string)
        print("Valid JSON Structure:", data)
        return data
    except json.JSONDecodeError as e:
        print("Invalid JSON Format:", str(e))
        raise ValueError(f"Invalid JSON: {str(e)}")
    
def sanitize_text(text):
    """
    Sanitize text by decoding HTML entities and removing special characters.

    Args:
        text (str): Input text to sanitize.

    Returns:
        str: Sanitized text.
    """
    text = html.unescape(text)
    text = re.sub(r'[^\w\s]', '', text)
    return text.strip()

def create_labels_pdf( labels_data, label_type):
    """
    Generate a PDF with labels based on the specified label type and product data.

    Args:
        labels_data (list): List of product details for the labels.
        label_type (str): Type of label (e.g., S-16987).

    Returns:
        str: Path to the generated PDF file.
    """
    try:
        if label_type not in LABEL_DIMENSIONS:
            raise ValueError(f"Unsupported label type: {label_type}")

        # Extract configuration for the label type
        config = LABEL_DIMENSIONS[label_type]
        label_width = config["label_width"]
        label_height = config["label_height"]
        labels_per_row = config["labels_per_row"]
        labels_per_column = config["labels_per_column"]
        margin_top = config["margin_top"] * 72
        margin_bottom = config["margin_bottom"] * 72
        margin_left = config["margin_left"] * 72
        margin_right = config["margin_right"] * 72
        file_name=config["file_name"]
        # Convert page size to points
        page_width_pts = config["page_width_inch"] * 72
        page_height_pts = config["page_height_inch"] * 72

        # Calculate usable area
        usable_width = page_width_pts - margin_left - margin_right
        usable_height = page_height_pts - margin_top - margin_bottom

        # Calculate spacing
        horizontal_spacing = (
            (usable_width - (labels_per_row * label_width * 72)) / (labels_per_row - 1)
            if labels_per_row > 1
            else 0
        )
        vertical_spacing = (
            (usable_height - (labels_per_column * label_height * 72)) / (labels_per_column - 1)
            if labels_per_column > 1
            else 0
        )

       # Get current date in YYYYMMDD format
        current_date = datetime.now().strftime('%Y%m%d')

        # Generate the output file path with the new naming convention
        output_path = os.path.join(
            os.path.abspath(DOC_FOLDER),
            f"{current_date}_{file_name}.pdf"
        )

        # Create canvas
        c = canvas.Canvas(output_path, pagesize=(page_width_pts, page_height_pts))

        # Initial offsets
        x_start = margin_left
        y_start = page_height_pts - margin_top
        x_offset = x_start
        y_offset = y_start

        label_count = 0

        for item in labels_data:
            sku = item["sku"]
            product = item["product"]
            price = item["display_price"]
            quantity = item["quantity"]

            for _ in range(quantity):
                # Start new page if needed
                if label_count > 0 and label_count % (labels_per_row * labels_per_column) == 0:
                    c.showPage()
                    x_offset = x_start
                    y_offset = y_start

                config = LABEL_DIMENSIONS[label_type]
                draw_label(c, x_offset, y_offset, sku, product, price, label_width, label_height, config)

                # Move to the next position
                x_offset += label_width * 72 + horizontal_spacing
                if (label_count + 1) % labels_per_row == 0:
                    x_offset = x_start
                    y_offset -= label_height * 72 + vertical_spacing

                label_count += 1

        # Save the PDF
        c.save()
        return output_path

    except Exception as e:
        raise Exception(f"Error in create_labels_pdf: {e}")

def draw_label(c, x, y, sku, name, price, label_width, label_height, config):
    """
    Draw a single label with QR code, product details, and wrapped text on the PDF canvas.

    Args:
        c (Canvas): ReportLab canvas object.
        x (float): X-coordinate of the label (in points).
        y (float): Y-coordinate of the label (in points).
        sku (str): SKU of the product.
        name (str): Product name.
        price (float): Product price.
        label_width (float): Width of the label (in inches).
        label_height (float): Height of the label (in inches).
        config (dict): Configuration for the label type including offsets.
    """
    # Convert inches to points
    label_width_pts = label_width * 72
    label_height_pts = label_height * 72
    qr_size_pts = min(label_width_pts, label_height_pts * 0.4)  # QR size is 40% of label height
    padding_pts = 4

    # Extract offsets from config
    qr_x_offset = config.get("qrcode_x_offset", 0) * 72
    qr_y_offset = -1 * config.get("qrcode_y_offset", 0) * 72
    sku_x_offset = config.get("sku_x_offset", 0) * 72
    sku_y_offset = -1 * config.get("sku_y_offset", 0) * 72
    price_x_offset = config.get("price_x_offset", 0) * 72
    price_y_offset = -1 * config.get("price_y_offset", 0) * 72
    show_product_name = config.get("show_product_name", False)

    qr_path = os.path.join(QR_FOLDER, f"{sku}.png")

    # Generate QR code if it doesn't exist
    if not os.path.exists(qr_path):
        qr = qrcode.QRCode(box_size=10, border=1)
        qr.add_data(sku)
        qr.make(fit=True)
        qr_img = qr.make_image(fill="black", back_color="white")
        qr_img.save(qr_path)

    # Position and draw the QR code
    c.drawImage(
        qr_path,
        x + (label_width_pts - qr_size_pts) / 2 + qr_x_offset,  # Center horizontally with offset
        y - qr_size_pts - padding_pts + qr_y_offset,  # Top position with offset
        width=qr_size_pts,
        height=qr_size_pts,
    )

    # Calculate the position for the SKU directly below the QR code
    sku_text_y = y - qr_size_pts - (5 * padding_pts)

    # Adjust SKU font size based on length
    max_font_size = 8
    min_font_size = 5
    styles = getSampleStyleSheet()
    text_style = styles["BodyText"]
    text_style.alignment = 1  # Center alignment

    # Adjust font size dynamically
    font_size = max_font_size
    while font_size >= min_font_size:
        text_style.fontSize = font_size
        text_style.leading = font_size + 2  # Adjust line spacing
        sku_paragraph = Paragraph(sku, text_style)
        sku_width = label_width_pts - (2 * padding_pts)  # Allow padding
        width, height = sku_paragraph.wrap(sku_width, qr_size_pts * 0.3)
        if height <= qr_size_pts * 0.3:
            break
        font_size -= 1

    sku_paragraph.drawOn(
        c,
        x + padding_pts,  # X-offset for padding
        sku_text_y + sku_y_offset  # Adjust Y position for SKU
    )

    if show_product_name:
        # Draw the product name below the SKU
        product_name_paragraph = Paragraph(name, text_style)
        product_name_width = label_width_pts - (2 * padding_pts)  # Allow padding
        product_name_paragraph.wrap(product_name_width, label_height_pts * 0.2)  # Adjust height
        product_name_paragraph.drawOn(
            c,
            x + padding_pts,
            sku_text_y - (2 * padding_pts)  # Adjust Y position for product name
        )

    # Draw the price at the bottom
    c.setFont("Helvetica-Bold", 10)
    c.drawCentredString(
        x + label_width_pts / 2 + price_x_offset,
        y - label_height_pts + (2 * padding_pts) + price_y_offset,  # Positioned above the bottom margin with offset
        f"${float(price):.2f}"
    )


@app.route('/upload', methods=['POST'])
def upload_and_process():
    """
    Handle file upload, validate, and preprocess data for label generation.

    Returns:
        Response: Rendered preview page or validation error messages.
    """
    # Dynamically reload label dimensions
    global LABEL_DIMENSIONS
    LABEL_DIMENSIONS = load_label_dimensions(LABEL_JSON_PATH)

    files = request.files.getlist('files[]')
    invalid_files = []
    valid_files = []
    aggregated_content = {}  # Dictionary to aggregate content by SKU
    total_labels=0
    for file in files:
        if not file.filename.endswith('.csv'):
            invalid_files.append({"file": file.filename, "reason": "Not a CSV file"})
            continue

        upload_path = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(upload_path)

        try:
            with open(upload_path, 'r') as csv_file:
                reader = csv.reader(csv_file)
                header = next(reader, None)

                if not header:
                    invalid_files.append({"file": file.filename, "reason": "Empty or invalid file"})
                    continue

                # Convert header to a mapping of column names to indices
                header_map = {col.strip(): idx for idx, col in enumerate(header)}

                if {'product', 'sku', 'quantity', 'display_price'}.issubset(header_map):
                    # Process File Type 1
                    for row in reader:
                        try:
                            sku = row[header_map['sku']]
                            product = sanitize_text(row[header_map['product']])
                            display_price = "{:.2f}".format(float(row[header_map['display_price']]))
                            quantity = int(row[header_map['quantity']])

                            if sku not in aggregated_content:
                                aggregated_content[sku] = {
                                    "sku": sku,
                                    "product": product,
                                    "display_price": display_price,
                                    "quantity": 0
                                }
                            aggregated_content[sku]["quantity"] += quantity

                        except ValueError as e:
                            invalid_files.append({"file": file.filename, "reason": f"Invalid quantity value in row: {row}"})


                elif {'name', 'sku', 'retail_price'}.issubset(header_map):
                    # Process File Type 2
                    inventory_columns = [col for col in header if re.match(r'^inventory_.*', col)]
                    if not inventory_columns:
                        invalid_files.append({"file": file.filename, "reason": "Missing inventory_* columns"})
                        continue

                    for row in reader:
                        try:
                            sku = row[header_map['sku']]
                            product_name = sanitize_text(" ".join([
                                row[header_map['name']],
                                row[header_map.get('variant_option_one_value', "")],
                                row[header_map.get('variant_option_two_value', "")],
                                row[header_map.get('variant_option_three_value', "")]
                            ]).strip())
                            display_price = "{:.2f}".format(float(row[header_map['retail_price']]))
                            quantity = sum(
                                int(row[header_map[col]]) for col in inventory_columns if row[header_map[col]].isdigit()
                            )

                            if sku not in aggregated_content:
                                aggregated_content[sku] = {
                                    "sku": sku,
                                    "product": product_name,
                                    "display_price": display_price,
                                    "quantity": 0
                                }
                            aggregated_content[sku]["quantity"] += quantity
                            total_labels += quantity
                            
                        except ValueError as e:
                            invalid_files.append({"file": file.filename, "reason": f"Invalid inventory value in row: {row}"})


        except Exception as e:
            invalid_files.append({"file": file.filename, "reason": f"Error reading file: {str(e)}"})

    if invalid_files:
        return jsonify({"message": "Some files failed validation", "invalid_files": invalid_files, "valid_files": valid_files}), 400

    if aggregated_content:
        # Convert aggregated content dictionary to a list for rendering
        processed_content = list(aggregated_content.values())
        return render_template("preview.html", processed_content=processed_content, label_types=LABEL_DIMENSIONS,total_labels=total_labels)

    return jsonify({"message": "No valid data to process", "valid_files": valid_files}), 400


@app.route('/generate_labels', methods=['POST'])
def generate_labels():
    """
    Generate labels in either PDF or Word format.
    """
    try:
        label_type = request.form.get("label_type")
        output_format = request.form.get("output_format", "pdf")
        processed_content = request.form.get("processed_content")

        if not label_type or label_type not in LABEL_DIMENSIONS:
            return jsonify({"message": "Invalid or missing label type"}), 400

        if not processed_content:
            raise ValueError("Processed content is empty")

        # Parse processed_content safely
        try:
            processed_content = json.loads(processed_content)
        except json.JSONDecodeError as e:
            return jsonify({"message": "Invalid processed content format", "error": str(e)}), 400

        # Get label configuration
        config = LABEL_DIMENSIONS[label_type]

        # Generate output based on format
        if output_format == "pdf":
            pdf_path = create_labels_pdf( processed_content, label_type)
            return send_file(pdf_path, as_attachment=True)
        elif output_format == "word":
            word_path = create_labels_word( processed_content, label_type, config)
            return send_file(word_path, as_attachment=True)
        else:
            return jsonify({"message": "Unsupported output format"}), 400

    except Exception as e:
        return jsonify({"message": "Error generating labels", "error": str(e)}), 500

def create_labels_word(file_code, labels_data, label_type, config):
    """
    Generate a Word document with labels based on the specified label type and product data.

    Args:
        file_code (str): Identifier for the file.
        labels_data (list): List of product details for the labels.
        label_type (str): Type of label (e.g., S-16987).
        config (dict): Configuration for the label type.

    Returns:
        str: Path to the generated Word file.
    """
    try:
        # Create a new Word document
        doc = Document()

        # Extract label configuration
        labels_per_row = config["labels_per_row"]
        labels_per_column = config["labels_per_column"]
        label_width_inch = config["label_width"]
        label_height_inch = config["label_height"]
        page_width_inch = config["page_width_inch"]
        page_height_inch = config["page_height_inch"]
        margin_top_inch = config["margin_top"]
        margin_left_inch = config["margin_left"]

        # Adjust page margins
        section = doc.sections[0]
        section.page_width = Inches(page_width_inch)
        section.page_height = Inches(page_height_inch)
        section.top_margin = Inches(margin_top_inch)
        section.bottom_margin = Inches(margin_top_inch)
        section.left_margin = Inches(margin_left_inch)
        section.right_margin = Inches(margin_left_inch)

        # Create a table to hold labels
        table = doc.add_table(rows=labels_per_column, cols=labels_per_row)
        table.allow_autofit = False

        # Set column widths
        for col in table.columns:
            col.width = Inches(label_width_inch)

        # Add labels to the table
        label_count = 0
        for item in labels_data:
            sku = item["sku"]
            product = item["product"]
            price = item["display_price"]
            quantity = item["quantity"]

            for _ in range(quantity):
                row_index = label_count // labels_per_row
                col_index = label_count % labels_per_row

                # Start a new page if needed
                if label_count > 0 and label_count % (labels_per_row * labels_per_column) == 0:
                    doc.add_page_break()
                    table = doc.add_table(rows=labels_per_column, cols=labels_per_row)
                    table.allow_autofit = False
                    for col in table.columns:
                        col.width = Inches(label_width_inch)

                # Get the cell for the current label
                cell = table.cell(row_index, col_index)

                # Add QR code
                qr_path = os.path.join("qr_codes", f"{sku}.png")
                if not os.path.exists(qr_path):
                    qr = qrcode.QRCode(box_size=10, border=1)
                    qr.add_data(sku)
                    qr.make(fit=True)
                    qr_img = qr.make_image(fill="black", back_color="white")
                    qr_img.save(qr_path)

                qr_paragraph = cell.add_paragraph()
                qr_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                qr_paragraph.add_run().add_picture(qr_path, width=Inches(label_width_inch * 0.4), height=Inches(label_height_inch * 0.4))

                # Add SKU
                sku_paragraph = cell.add_paragraph()
                sku_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sku_run = sku_paragraph.add_run(f"SKU: {sku}")
                sku_run.bold = True
                sku_run.font.size = Pt(8)

                # Add product name
                product_paragraph = cell.add_paragraph()
                product_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                product_run = product_paragraph.add_run(product)
                product_run.font.size = Pt(8)

                # Add price
                price_paragraph = cell.add_paragraph()
                price_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                price_run = price_paragraph.add_run(f"Price: ${float(price):.2f}")
                price_run.bold = True
                price_run.font.size = Pt(8)

                label_count += 1

        # Ensure the directory exists
        word_folder = os.path.join(os.getcwd(), "word_ready_files")
        os.makedirs(word_folder, exist_ok=True)

        # Save the Word file
        word_path = os.path.join(word_folder, f"Labels_{file_code}_{label_type}.docx")
        doc.save(word_path)
        return word_path

    except Exception as e:
        raise Exception(f"Error in create_labels_word: {e}")

@app.route('/')
def home():
    """
    Render the home page for file uploads.

    Returns:
        Response: Rendered HTML page.
    """
    return render_template('index.html')

if __name__ == '__main__':
    app.run(debug=True)
